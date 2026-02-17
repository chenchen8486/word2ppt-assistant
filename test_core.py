import os
import docx
from core.doc_loader import DocLoader
from core.ppt_generator import PPTGenerator
from utils.logger import setup_logger

# Setup logger
setup_logger()

def create_dummy_docx(path):
    doc = docx.Document()
    doc.add_heading('Test Exam', 0)
    doc.add_paragraph('This is a test context paragraph.')
    doc.add_paragraph('1. What is Python?')
    doc.add_paragraph('A. A snake')
    doc.add_paragraph('B. A programming language')
    doc.save(path)
    print(f"Created dummy doc at {path}")

def test_workflow():
    docx_path = "test_data/test_exam.docx"
    pptx_path = "test_data/test_exam_output.pptx"
    
    # 1. Create Doc
    if not os.path.exists("test_data"):
        os.makedirs("test_data")
    create_dummy_docx(docx_path)
    
    # 2. Test Loader
    loader = DocLoader(docx_path)
    loader.load()
    chunks = list(loader.get_chunks())
    print(f"Chunks generated: {len(chunks)}")
    assert len(chunks) > 0
    
    # 3. Mock LLM Data
    mock_questions = [
        {
            "id": "1",
            "type": "Single Choice",
            "context": "This is a test context paragraph.",
            "question": "What is Python?",
            "options": ["A. A snake", "B. A programming language"],
            "answer": "B",
            "analysis": "Python is a high-level programming language."
        }
    ]
    
    # 4. Test PPT Generator
    gen = PPTGenerator()
    gen.generate(mock_questions, pptx_path, title="Test Exam")
    
    if os.path.exists(pptx_path):
        print("PPT generated successfully!")
    else:
        print("PPT generation failed.")

if __name__ == "__main__":
    test_workflow()
