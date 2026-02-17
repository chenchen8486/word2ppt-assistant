import os
import re
import docx
from utils.logger import get_logger

logger = get_logger()

class DocLoader:
    """
    Handles loading and structural chunking of Word documents (.docx).
    """
    def __init__(self, file_path):
        self.file_path = file_path
        self.paragraphs = []
        self.full_text = ""
        self.total_questions_detected = 0 # Watchdog counter

    def load(self):
        """
        Loads the .docx file and extracts text from paragraphs.
        """
        if not os.path.exists(self.file_path):
            error_msg = f"File not found: {self.file_path}"
            logger.error(error_msg)
            raise FileNotFoundError(error_msg)

        try:
            logger.info(f"Loading document: {self.file_path}")
            doc = docx.Document(self.file_path)
            
            # Extract text from paragraphs, skipping empty ones
            self.paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
            self.full_text = "\n".join(self.paragraphs)
            
            # Pre-scan for question numbers (Watchdog Input Counter)
            self._count_questions()
            
            logger.info(f"Successfully loaded document. Total paragraphs: {len(self.paragraphs)}")
            return True
        except Exception as e:
            logger.error(f"Failed to load document: {e}")
            raise

    def _count_questions(self):
        """
        Estimates total number of questions using Regex.
        Matches patterns like: "1.", "2.", "Question 1", "(1)", "【1】"
        """
        # Regex patterns for common question starts
        patterns = [
            r'^\d+\.',           # 1.
            r'^Question\s+\d+',  # Question 1
            r'^\(\d+\)',         # (1)
            r'^【\d+】',         # 【1】
            r'^第\d+题'          # 第1题
        ]
        
        # Pattern for Big Question Titles (e.g. "一、选择题 (30分)")
        # We want to ignore these when counting actual questions
        big_question_pattern = r'^[一二三四五六七八九十]+、'

        count = 0
        for p in self.paragraphs:
            # Skip Big Question Titles
            if re.match(big_question_pattern, p):
                continue
                
            for pat in patterns:
                if re.match(pat, p):
                    count += 1
                    break
        
        self.total_questions_detected = count
        logger.info(f"[Watchdog] Estimated total questions in doc: {count}")

    def get_chunks(self, max_chars=3000):
        """
        Yields structural chunks based on question boundaries.
        Prevents splitting a single question into two chunks.
        """
        if not self.paragraphs:
            self.load()

        current_chunk = []
        current_length = 0
        chunk_index = 1
        
        # Regex to detect start of a new question
        question_start_pattern = re.compile(r'^\d+\.|^Question\s+\d+|^\(\d+\)|^【\d+】|^第\d+题')
        
        # Regex for Big Question Titles (Context)
        # e.g., "一、阅读理解(20分)", "二、选择题"
        big_context_pattern = re.compile(r'^[一二三四五六七八九十]+、')

        i = 0
        while i < len(self.paragraphs):
            para = self.paragraphs[i]
            para_len = len(para)
            is_new_question = bool(question_start_pattern.match(para))
            is_big_context = bool(big_context_pattern.match(para))

            # Strategy:
            # 1. If it's a Big Context (e.g. "一、..."), it usually starts a new section.
            #    We should probably start a new chunk to keep context with its questions,
            #    UNLESS the current chunk is very small.
            # 2. If it's a new question ("1."), we check length.
            
            should_split = False
            
            # Priority Split: Big Context
            # If we hit "二、...", and we have pending questions, split now.
            if is_big_context and current_chunk:
                should_split = True
            
            # Regular Split: Length limit reached + Safe break point
            elif current_length > max_chars:
                if is_new_question:
                    should_split = True
                elif current_length > max_chars * 1.5:
                    # Forced split protection
                    should_split = True
                    logger.warning(f"Chunk #{chunk_index} forced split due to excessive length (no question break found).")

            if should_split and current_chunk:
                chunk_text = "\n".join(current_chunk)
                logger.info(f"Generated Structural Chunk #{chunk_index} (Length: {len(chunk_text)})")
                yield chunk_text
                
                chunk_index += 1
                current_chunk = []
                current_length = 0
            
            current_chunk.append(para)
            current_length += para_len
            i += 1
            
        # Yield the last chunk
        if current_chunk:
            chunk_text = "\n".join(current_chunk)
            logger.info(f"Generated Structural Chunk #{chunk_index} (Length: {len(chunk_text)})")
            yield chunk_text

if __name__ == "__main__":
    # Test block
    pass
