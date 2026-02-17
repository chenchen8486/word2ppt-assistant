import os
import re
import docx
from docx.document import Document
from utils.logger import get_logger

logger = get_logger()

class Standardizer:
    def __init__(self, file_path, output_image_dir="images", llm_client=None):
        self.file_path = file_path
        self.output_image_dir = output_image_dir
        self.llm_client = llm_client
        self.doc = None
        self.paragraphs = [] # List of {'text': str, 'images': [], 'original_index': int}
        self.questions = {}  # Map id (str) -> dict
        self.analysis_map = {} # Map id (str) -> {'answer': str, 'analysis': str}
        
        # Regex Patterns
        self.start_pattern = re.compile(r'^[一二三四五六七八九十]+、') 
        self.small_question_pattern = re.compile(r'^(\d+)\.')
        # Matches 【1题详解】 or 1题详解
        self.analysis_header_pattern = re.compile(r'^(?:【)?(\d+)题详解(?:】)?')
        # Matches 【答案】
        self.answer_header_pattern = re.compile(r'^【答案】')
        
        self.exclude_patterns = [
            re.compile(r'^【导语】'),
            re.compile(r'^参考译文')
        ]

        if not os.path.exists(self.output_image_dir):
            os.makedirs(self.output_image_dir)

    def load_and_parse(self):
        """
        Main entry point.
        """
        if not os.path.exists(self.file_path):
            raise FileNotFoundError(f"File not found: {self.file_path}")

        logger.info(f"Loading document: {self.file_path}")
        self.doc = docx.Document(self.file_path)
        
        # 1. Extract raw content
        self._extract_content()
        
        # 2. Filter content (Start point, Excludes)
        self._filter_content()
        
        # 3. Two-Pass Parsing
        self._pass_1_extract_questions()
        self._pass_2_extract_analysis()
        
        # 4. Merge
        self._merge_data()
        
        # 5. LLM Refinement (Options extraction)
        if self.llm_client:
            self._refine_questions_with_llm()
        
        # 6. Build Result
        return self._build_result_list()

    def _extract_content(self):
        logger.info("Extracting content and images...")
        image_count = 0
        self.paragraphs = []
        
        for i, para in enumerate(self.doc.paragraphs):
            text = para.text.strip()
            # Basic cleanup
            text = text.replace('\xa0', ' ') 
            
            images = []
            # Extract images from runs
            for run in para.runs:
                if 'drawing' in run._element.xml:
                    drawings = run._element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing')
                    for drawing in drawings:
                        blips = drawing.findall('.//{http://schemas.openxmlformats.org/drawingml/2006/main}blip')
                        for blip in blips:
                            embed_id = blip.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
                            if embed_id:
                                try:
                                    image_part = self.doc.part.related_parts[embed_id]
                                    content_type = image_part.content_type
                                    ext = 'jpg'
                                    if 'png' in content_type: ext = 'png'
                                    elif 'jpeg' in content_type: ext = 'jpg'
                                    
                                    filename = f"img_{i}_{image_count}.{ext}"
                                    path = os.path.join(self.output_image_dir, filename)
                                    
                                    with open(path, "wb") as f:
                                        f.write(image_part.blob)
                                    
                                    images.append(path)
                                    image_count += 1
                                except Exception as e:
                                    logger.warning(f"Image extraction failed: {e}")

            self.paragraphs.append({
                'text': text,
                'images': images,
                'original_index': i
            })

    def _filter_content(self):
        start_index = -1
        # Look for "一、（x分）"
        # User said: "一、（9分）"
        # My regex: ^[一二三四五六七八九十]+、
        # Let's find the first occurrence of this pattern.
        
        for i, p in enumerate(self.paragraphs):
            if self.start_pattern.match(p['text']):
                start_index = i
                logger.info(f"Start point found at line {i}: {p['text']}")
                break
        
        if start_index == -1:
            logger.warning("Start pattern not found. Using all content.")
            start_index = 0
            
        filtered = []
        for i in range(start_index, len(self.paragraphs)):
            p = self.paragraphs[i]
            text = p['text']
            should_exclude = any(pat.search(text) for pat in self.exclude_patterns)
            if not should_exclude:
                filtered.append(p)
                
        self.paragraphs = filtered

    def _pass_1_extract_questions(self):
        """
        Identify Contexts and Questions.
        We stop treating things as questions if we suspect we entered the Analysis section globally?
        User said "Separated structure".
        But regex is safer.
        """
        current_context = ""
        current_q_id = None
        current_lines = []
        current_images = []
        
        for p in self.paragraphs:
            text = p['text']
            
            # 1. Check for Analysis Header -> If found, we might be in Analysis land.
            # But in Pass 1, we just want to NOT classify analysis as questions.
            # If line matches "1题详解", it's definitely not a question start.
            if self.analysis_header_pattern.match(text):
                # End current question
                if current_q_id:
                    self._save_question(current_q_id, current_context, current_lines, current_images)
                    current_q_id = None
                    current_lines = []
                    current_images = []
                continue # Skip analysis lines in Pass 1

            # 2. Check for Context (Big Question)
            if self.start_pattern.match(text):
                if current_q_id:
                    self._save_question(current_q_id, current_context, current_lines, current_images)
                    current_q_id = None
                    current_lines = []
                    current_images = []
                current_context = text
                continue

            # 3. Check for Small Question Start
            match = self.small_question_pattern.match(text)
            if match:
                # New Question
                if current_q_id:
                    self._save_question(current_q_id, current_context, current_lines, current_images)
                
                current_q_id = match.group(1)
                current_lines = [text]
                current_images = p['images']
            else:
                # Continuation
                if current_q_id:
                    # Check if this line looks like an Answer line?
                    # If "【答案】" appears, it might be the end of the question text?
                    if self.answer_header_pattern.match(text):
                        # End of question text
                        self._save_question(current_q_id, current_context, current_lines, current_images)
                        current_q_id = None
                        current_lines = []
                        current_images = []
                    else:
                        current_lines.append(text)
                        current_images.extend(p['images'])
                elif current_context:
                    # Append to context if no question active
                    # But be careful not to append analysis text to context
                    if not self.answer_header_pattern.match(text):
                        current_context += "\n" + text

        # Save last
        if current_q_id:
            self._save_question(current_q_id, current_context, current_lines, current_images)

    def _save_question(self, q_id, context, lines, images):
        if q_id in self.questions:
            # Duplicate ID? 
            # In exam papers, sometimes IDs restart?
            # Assuming unique IDs for now as per user instruction "题号关联".
            pass
        
        full_text = "\n".join(lines).strip()
        self.questions[q_id] = {
            "id": q_id,
            "type": "question",
            "context": context,
            "question": full_text, # To be refined
            "options": [],
            "answer": "",
            "analysis": "",
            "images": images
        }

    def _pass_2_extract_analysis(self):
        """
        Identify Answers and Analysis.
        """
        current_ana_id = None
        
        for p in self.paragraphs:
            text = p['text']
            
            # Check for Analysis Header
            match = self.analysis_header_pattern.match(text)
            if match:
                current_ana_id = match.group(1)
                if current_ana_id not in self.analysis_map:
                    self.analysis_map[current_ana_id] = {'answer': '', 'analysis': ''}
                
                # Append header line to analysis? User says "ensure captured".
                self.analysis_map[current_ana_id]['analysis'] += text + "\n"
                continue
                
            # Check for Answer Header
            match_ans = self.answer_header_pattern.match(text)
            if match_ans:
                # Extract answer content
                # Format: "【答案】 A" or "【答案】\n A"
                content = text.replace("【答案】", "").strip()
                
                # Assign to WHO?
                # If we are inside an Analysis Block (current_ana_id is set), assign to it.
                if current_ana_id:
                    self.analysis_map[current_ana_id]['answer'] = content
                else:
                    # "散落在各处" -> Maybe it's just before/after something?
                    # If we are NOT in analysis block, maybe we can't link it easily unless it has an ID nearby?
                    # BUT, usually "【答案】" is followed by "1. A 2. B"?
                    # Or "【答案】" is inside the question block (which we handled in Pass 1 to stop question)?
                    # If it's a standalone "【答案】" line with no ID, it's hard.
                    # Let's assume strict structure: "【1题详解】...【答案】A..."
                    pass
                continue

            # Continuation
            if current_ana_id:
                # If we hit a new Question or Context, we stop analysis?
                # But Pass 1 already filtered those.
                # Just check if we hit another header.
                if self.start_pattern.match(text) or self.small_question_pattern.match(text):
                    current_ana_id = None # Reset
                else:
                    self.analysis_map[current_ana_id]['analysis'] += text + "\n"

    def _merge_data(self):
        for q_id, q_data in self.questions.items():
            if q_id in self.analysis_map:
                ana = self.analysis_map[q_id]
                q_data['answer'] = ana['answer']
                q_data['analysis'] = ana['analysis']
            else:
                logger.warning(f"Missing analysis for Question {q_id}")
                q_data['analysis'] = "【解析待补充】"

    def _refine_questions_with_llm(self):
        """
        Uses LLM to split 'question' text into 'stem' and 'options'.
        Uses concurrent processing for speed.
        """
        from concurrent.futures import ThreadPoolExecutor, as_completed
        
        logger.info("Refining questions with LLM...")
        
        q_ids = sorted(self.questions.keys(), key=lambda x: int(x) if x.isdigit() else 999)
        batch_size = 5
        batches = [q_ids[i:i+batch_size] for i in range(0, len(q_ids), batch_size)]
        
        total_batches = len(batches)
        completed_batches = 0
        
        def process_batch(batch_ids):
            batch_text = ""
            for qid in batch_ids:
                batch_text += f"--- Question {qid} ---\n{self.questions[qid]['question']}\n\n"
            
            sys_prompt = (
                "You are an exam question parser. Your task is to split raw question text into 'Question Stem' and 'Options'.\n"
                "Return a JSON LIST of objects. Each object must have:\n"
                "- `id`: The question number (string).\n"
                "- `question`: The stem of the question.\n"
                "- `options`: A list of strings for options (e.g. ['A. xxx', 'B. xxx']). If no options, return empty list.\n"
                "Strictly return valid JSON."
            )

            try:
                result = self.llm_client.process_chunk(
                    batch_text, 
                    system_prompt_override=sys_prompt
                )
                return result
            except Exception as e:
                logger.error(f"LLM refinement failed: {e}")
                return None

        # Use max 5 workers
        with ThreadPoolExecutor(max_workers=5) as executor:
            future_to_batch = {executor.submit(process_batch, batch): batch for batch in batches}
            
            for future in as_completed(future_to_batch):
                result = future.result()
                if result:
                    for item in result:
                        qid = str(item.get('id'))
                        if qid in self.questions:
                            self.questions[qid]['question'] = item.get('question', '')
                            self.questions[qid]['options'] = item.get('options', [])
                
                completed_batches += 1
                logger.info(f"Refinement progress: {completed_batches}/{total_batches} batches")

    def _build_result_list(self):
        # Return flat list with Context objects injected?
        # User wants "List of Dict... id, title, options, answer, analysis, image".
        # And "Context" logic handled by PPTGenerator?
        # PPTGenerator handles "context" field in Question object.
        # But wait, earlier I changed PPTGenerator to handle "type": "context" objects.
        # "Standardizer" should probably output that format.
        
        final_list = []
        
        # Group by Context
        # Since my Questions map stores context string, I can regroup.
        # Or just emit questions and let PPTGenerator handle "Background (Question X)".
        # BUT, PPTGenerator's new logic (sequential) expects Context Object -> Question Object.
        
        # Sort questions
        sorted_qs = sorted(self.questions.values(), key=lambda x: int(x['id']) if x['id'].isdigit() else 9999)
        
        last_context = None
        for q in sorted_qs:
            ctx = q['context']
            if ctx and ctx != last_context:
                final_list.append({
                    "type": "context",
                    "content": ctx
                })
                last_context = ctx
            
            final_list.append(q)
            
        return final_list
